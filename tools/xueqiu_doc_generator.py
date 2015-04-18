import urllib2

from BeautifulSoup import BeautifulSoup, Comment
from docx import Document
from docx.shared import Inches
import StringIO
import json
from utils import soup_wrapper


req_header={'User-Agent':'Mozilla/5.0 (Windows NT 6.1) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/38.0.2125.122 Safari/537.36',
'Accept':'text/html;q=0.9,*/*;q=0.8',
'Accept-Charset':'ISO-8859-1,utf-8;q=0.7,*;q=0.3',
'Connection':'keep-alive',
 'Host': 'xueqiu.com',
 'Cookie':'bid=33fc33407174511da049cc98ca6f6ad6_i7yxho92; xq_a_token=c5e743b95431d7c27af48e78e042a68af4b2627f; xqat=c5e743b95431d7c27af48e78e042a68af4b2627f; xq_r_token=98b0991c0da598599c85c54bee7ccb0f6c6c5578; xq_is_login=1; xq_token_expire=Mon%20Apr%2027%202015%2022%3A20%3A00%20GMT%2B0800%20(CST); __utma=1.247665118.1427904398.1428742729.1428804366.11; __utmc=1; __utmz=1.1428073731.4.2.utmcsr=haosou.com|utmccn=(organic)|utmcmd=organic|utmctr=xue; Hm_lvt_1db88642e346389874251b5a1eded6e3=1428073731,1428320714,1428721390,1428804366; Hm_lpvt_1db88642e346389874251b5a1eded6e3=1428804778'
}

def findLinksAndGenDocs(authors):
    for author in authors :
        print author["name"]
       
        url = author["link"]
        retText = soup_wrapper.getPageContent(url, req_header)
        jsonObject = json.loads(retText);
        maxPage = jsonObject["maxPage"]
        print maxPage
        
        count = 0
        for i in range(1, maxPage+1) :
            url = url.replace("page=1", "page=%d" %i)
            pageText = soup_wrapper.getPageContent(url, req_header)
            pageObject = json.loads(pageText);
            articles = pageObject["statuses"]
            for article in articles :
                articePage = "http://xueqiu.com/%s" %article["target"]
                print articePage
                parsePageAndGenDocFile(articePage)
        
        
        print count


def parsePageAndGenDocFile(link):
    cotent = soup_wrapper.getPageContent(link, req_header)
    if cotent == "" :
        return
    soupContent = BeautifulSoup(cotent)
    title = ""
    if soupContent.find(attrs={"class" : "status-title"}) :
        title = soupContent.find(attrs={"class" : "status-title"}).contents[0]
    print title
    pageContent = soupContent.find(attrs={"class" : "detail"})
    comments = soupContent.findAll(text=lambda text:isinstance(text, Comment))
    [comment.extract() for comment in comments]
    
    scripts = soupContent.findAll("script")
    [script.extract() for script in scripts]
    
    brs = soupContent.findAll("br")
    [br.extract() for br in brs]
    
    document = Document()
    document.add_heading(title, 0)
    
    for line in pageContent.contents :
        try :
            if  line.find("img class=\"ke_img\"") != -1 :
                imagePath = line.get("src").encode("utf-8").replace("!custom.jpg", "")
    #             imagePath = imagePath.findAll("^http:.*jpg")
                print imagePath
                
                image_from_url = urllib2.urlopen(imagePath)
                io_url = StringIO.StringIO()
                io_url.write(image_from_url.read())
                io_url.seek(0)
                document.add_picture(io_url, width=Inches(4))
            else :   
                document.add_paragraph(line)
        except Exception,ex:
            print line
    
    title = title.replace("-", "")
    title = title.replace("?", "")
    document.save(title + '.docx')


if __name__ == '__main__':
    
    authors = [
#             {'name' : "chaodeshixin",'link' :'http://xueqiu.com/v4/statuses/user_timeline.json?user_id=2821861040&page=1&type=2&_=1428814922049' }, 
#             {'name': "tangshizhuren", 'link': 'http://xueqiu.com/v4/statuses/user_timeline.json?user_id=2054435398&page=1&type=2&_=1428817401323'}, 
#             {'name' : "kezhong",'link' :'http://xueqiu.com/v4/statuses/user_timeline.json?user_id=5243796549&page=1&type=2&_=142881504540' }
            {'name': "mengting",'link': 'http://xueqiu.com/v4/statuses/user_timeline.json?user_id=1965949492&page=1&type=2&_=1428815110011'}
            ]
    findLinksAndGenDocs(authors)
    
