# -*- coding: utf-8 -*-
import re
import urllib2

from BeautifulSoup import BeautifulSoup, Comment
from docx import Document
from docx.shared import Inches
import StringIO


def findLinksAndGenDocs():
    page = getPageContent("http://xueqiu.com/2054435398/32283614")
    soup = BeautifulSoup(page)
    
    links = soup.find(attrs={"class" : "detail"}).findAll(attrs={"title" : re.compile("^http")})
    for link in links :
#         print link['href']
#         if link["href"] == "http://xueqiu.com/2054435398/31741295" :
        parsePageAndGenDocFile(link['href'])
#         break


def getPageContent(link):
    req_header={'User-Agent':'Mozilla/5.0 (Windows NT 6.1) AppleWebKit/537.11 (KHTML, like Gecko) Chrome/23.0.1271.64 Safari/537.11',
    'Accept':'text/html;q=0.9,*/*;q=0.8',
    'Accept-Charset':'ISO-8859-1,utf-8;q=0.7,*;q=0.3',
    'Connection':'close',
    'Referer':None
    }
    try :
        req_timeout=1000
        req=urllib2.Request(link, None, req_header)
        resp=urllib2.urlopen(req, None, req_timeout)
        html = resp.read()
    except Exception,ex:
        html = ""
    html = unicode(html, "utf8")
    return html

def parsePageAndGenDocFile(link):
    cotent = getPageContent(link)
    if cotent == "" :
        return
    soupContent = BeautifulSoup(cotent)
    title = soupContent.find(attrs={"class" : "status-title"}).contents[0]
    print title
    pageContent = soupContent.find(attrs={"class" : "detail"})
    comments = soupContent.findAll(text=lambda text:isinstance(text, Comment))
    [comment.extract() for comment in comments]
    
    scripts = soupContent.findAll("script")
    [script.extract() for script in scripts]
    
    brs = soupContent.findAll("br")
    [br.extract() for br in brs]
    
    document = Document()
    document.add_heading(title, 0)
    
    for line in pageContent.contents :
        try :
            if  line.find("img class=\"ke_img\"") != -1 :
                imagePath = line.get("src").encode("utf-8").replace("!custom.jpg", "")
    #             imagePath = imagePath.findAll("^http:.*jpg")
                print imagePath
                
                image_from_url = urllib2.urlopen(imagePath)
                io_url = StringIO.StringIO()
                io_url.write(image_from_url.read())
                io_url.seek(0)
                document.add_picture(io_url, width=Inches(4))
            else :   
                document.add_paragraph(line)
        except Exception,ex:
            print line
    
    title = title.replace("-", "")
    title = title.replace("?", "")
    document.save(title + '.docx')

if __name__ == '__main__':
    findLinksAndGenDocs()
    